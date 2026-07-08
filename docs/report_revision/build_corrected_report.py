from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "CyberSentinel_Final_Report_Original.docx"
OUTPUT = ROOT / "CyberSentinel_Final_Report_Corrected.docx"
DIAGRAMS = ROOT / "diagram_assets"


def set_paragraph_text(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.add_run(text)


def replace_caption_suffix(paragraph, suffix):
    if not paragraph.runs:
        paragraph.add_run(suffix)
        return
    paragraph.runs[-1].text = suffix


def set_cell(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)


def set_row(table, row_index, values):
    row = table.rows[row_index]
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def delete_row(table, row_index):
    row = table.rows[row_index]
    row._tr.getparent().remove(row._tr)


def set_page_numbering(section, fmt, start):
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:fmt"), fmt)
    node.set(qn("w:start"), str(start))


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    return run


def clear_story(story):
    for child in list(story._element):
        # Headers and footers can wrap existing content in structured-document
        # tags (w:sdt), including Word's built-in page-number control. Remove
        # every story child so rebuilding the footer cannot leave a duplicate
        # page field behind.
        story._element.remove(child)
    story.add_paragraph()


def format_run(run, size=9, italic=False, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    if run._element.get_or_add_rPr().rFonts is not None:
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")


def build_footer(section, include_page=True):
    clear_story(section.footer)
    paragraph = section.footer.paragraphs[0]
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.45), WD_TAB_ALIGNMENT.RIGHT
    )
    left = paragraph.add_run("Department of Computer Science, CUI, Abbottabad")
    format_run(left, 8, italic=True)
    if include_page:
        paragraph.add_run("\t")
        page_run = add_page_field(paragraph)
        format_run(page_run, 8, italic=True)


def build_header(section, text):
    clear_story(section.header)
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    format_run(run, 8, italic=True)


def add_front_body_section_break(document):
    body_heading = document.paragraphs[176]
    previous = document.paragraphs[175]
    p_pr = previous._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    copied = deepcopy(document.sections[-1]._sectPr)
    type_node = copied.find(qn("w:type"))
    if type_node is None:
        type_node = OxmlElement("w:type")
        copied.insert(0, type_node)
    type_node.set(qn("w:val"), "nextPage")
    p_pr.append(copied)
    set_paragraph_text(body_heading, "1. Introduction")
    direct_numbering = body_heading._p.pPr.find(qn("w:numPr"))
    if direct_numbering is not None:
        body_heading._p.pPr.remove(direct_numbering)


def add_front_matter_page_breaks(document):
    for index in [85, 102, 114, 145]:
        document.paragraphs[index].paragraph_format.page_break_before = True
    for paragraph in document.element.body.xpath(".//w:p"):
        text = "".join(paragraph.xpath(".//w:t/text()"))
        if text.strip().lower() == "table of contents":
            p_pr = paragraph.get_or_add_pPr()
            if p_pr.find(qn("w:pageBreakBefore")) is None:
                p_pr.append(OxmlElement("w:pageBreakBefore"))


PARAGRAPHS = {
    170: "",
    76: "Cybersecurity threats such as phishing emails, malicious URLs, and malware-infected files continue to affect individuals and small organizations. The risk is especially serious for non-technical users who may not have access to affordable security tools or the expertise needed to interpret conventional threat-intelligence reports.",
    77: "CyberSentinel AI - Intelligent Cybersecurity Toolkit for Phishing, Malware, and Threat Detection addresses this problem through a unified web application. It allows users to analyze email text and supported documents, inspect URLs, and submit files for malware reputation analysis through a simple dashboard.",
    78: "The email module applies Natural Language Processing with word-level and character-level TF-IDF features and a binary Logistic Regression classifier that predicts ham or spam. The interface maps ham to Safe and spam to Phishing. Visible HTTP/HTTPS links found in email text are analyzed separately. The URL module uses ten rule-based structural checks together with AbuseIPDB reputation data, while the malware module uses SHA-256 lookup and live analysis through VirusTotal [8], [14], [15].",
    79: "The implementation uses a React frontend, a Flask backend, and scikit-learn for the phishing model. PostgreSQL provides persistent storage in the deployed environment, while SQLite remains the local-development fallback. The backend is deployed on Render and the frontend on Vercel; the Gmail Chrome extension connects to the same backend [9], [11], [12], [24], [33], [34].",
    80: "Results are presented through color-coded verdicts, confidence values, threat indicators, and safety recommendations. Scan summaries are stored in the database, and PDF reports are generated in the browser from the current result. Email alerts are attempted only for configured high-risk outcomes and do not interrupt scanning when SMTP delivery fails.",
    81: "CyberSentinel AI therefore demonstrates a practical AI-assisted cybersecurity prototype. It combines a trained phishing classifier with deterministic URL rules and external file-reputation services while clearly separating implemented capabilities from future security hardening and advanced threat-analysis work.",
    179: "CyberSentinel AI is a web-based cybersecurity platform that combines three implemented detection approaches: machine-learning classification for phishing email text, rule-based analysis for suspicious URLs, and VirusTotal-assisted reputation analysis for uploaded files.",
    180: "Users can submit email text, supported email or document files, URLs, and malware-scan files. Email results are presented as Safe or Phishing, URL results as Safe, Suspicious, or Malicious, and file results using VirusTotal-derived verdicts such as Clean, Suspicious, Malicious, Unknown, or Pending.",
    181: "The phishing module uses hybrid TF-IDF features and Logistic Regression. The URL module is deterministic rather than machine-learning based and combines ten structural rules with AbuseIPDB IP reputation. The file module computes SHA-256 hashes and uses VirusTotal for known-file lookup and, when necessary, live upload analysis [14], [15].",
    186: "• Machine Learning: Implementation and evaluation of a Logistic Regression classifier for binary phishing-email detection.",
    190: "• Database Systems: Storage of users and email, URL, and file scan summaries in PostgreSQL, with SQLite used as a local fallback.",
    197: "CyberSentinel AI addresses these challenges through one web interface that combines an ML-based phishing classifier, rule-based URL scoring, and VirusTotal-assisted file analysis. This hybrid approach keeps each module aligned with the technique actually suited to its input.",
    201: "VirusTotal is widely used to inspect files and URLs using results contributed by multiple security engines. It is effective for known-file reputation but its detailed results can be difficult for non-technical users to interpret [14].",
    202: "PhishTank, Kaspersky Threat Intelligence Portal, and similar services provide useful phishing or malware intelligence, but they are separate platforms rather than a unified end-user workflow. PhishTank was reviewed during the project but is not integrated into the final implementation [16], [38].",
    203: "Cisco Talos Intelligence provides broad threat-intelligence and network-security information, but its professional focus can be difficult for general users who only need a clear decision and recommended action [37].",
    204: "Research and standard machine-learning libraries support classifiers such as Logistic Regression, Support Vector Machines, and transformer-based models for text classification. CyberSentinel AI selects Logistic Regression because it is fast, interpretable, and suitable for sparse TF-IDF features [4], [5], [6], [8].",
    209: "A single detection technique is not sufficient for every threat type. CyberSentinel AI therefore uses a trained text classifier for email content, deterministic structural rules for URLs, and continuously updated external reputation data for files and resolved IP addresses.",
    210: "The final system combines these techniques without claiming that every module uses machine learning. VirusTotal and AbuseIPDB provide live external intelligence, while the locally loaded phishing model provides low-latency text classification [14], [15].",
    237: "Integration with implemented external services: VirusTotal, AbuseIPDB, and SMTP email delivery",
    264: "The User registers or logs in, performs email, URL, and file scans, reviews the latest scan records, downloads result reports, and may scan an open Gmail message through the extension. The Administrator can view system statistics, registered users, and recent scan records; user editing, settings management, and model management are outside the current prototype.",
    292: "Rule-based analyzer evaluates URL features",
    293: "AbuseIPDB provides IP reputation when available",
    334: "CyberSentinel AI follows a modular client-server architecture comprising React and Chrome-extension clients, Flask API routes, detection services, persistent storage, and external intelligence services. The modules communicate through JSON or multipart HTTP requests.",
    335: "The client layer contains the React web application and the Gmail Chrome extension. The web interface provides email, URL, file, history, report, and administrative views, while the extension submits the visible text of an opened Gmail message.",
    336: "The Flask backend validates requests, applies rate limits and payload controls, invokes the appropriate detection service, returns JSON responses, stores scan summaries, and attempts configured email notifications. External integrations are limited to VirusTotal, AbuseIPDB, and SMTP [11], [14], [15], [26].",
    337: "The service layer contains the phishing NLP/ML pipeline, the deterministic URLAnalyzer, the MalwareScanner, database utilities, and notification functions. Only the phishing email classifier uses machine learning; URL decisions are rule-based and malware verdicts are derived from VirusTotal responses.",
    338: "The data layer uses PostgreSQL when DATABASE_URL is configured and SQLite otherwise. It stores users plus email_scans, url_scans, and file_scans records. Application log files and browser-generated PDF reports are not database entities.",
    343: "The Entity Relationship diagram represents the four tables created by the current database initialization code: users, email_scans, url_scans, and file_scans.",
    344: "Each scan table stores a user_email value together with module-specific result fields and a timestamp. These values provide a logical relationship to users.email, although the current schema does not enforce foreign-key constraints.",
    345: "PostgreSQL is used for persistent live deployment and SQLite is retained for local development. Both database paths expose the same logical table structure.",
    351: "The process flow diagram describes the implemented path from user input through validation, module-specific analysis, database storage, result presentation, and optional client-side PDF generation.",
    352: "Email text and extracted document text pass through sanitization, hybrid TF-IDF transformation, and Logistic Regression; up to five unique visible URLs may also be analyzed. Submitted URLs use structural scoring and AbuseIPDB. Uploaded files use SHA-256 lookup and optional VirusTotal live upload.",
    353: "After a scan, a summary record is stored in the appropriate scan table and the JSON result is displayed with verdicts, confidence values, indicators, and recommendations. The user can then download the current result as a PDF or open the latest-history view.",
    357: "4.4 Class and Module Diagram",
    358: "The implementation is organized around Flask blueprints, two service classes, and supporting modules rather than separate domain classes for User, Admin, ScanSession, or Report.",
    359: "URLAnalyzer and MalwareScanner are the principal service classes. Prediction, notification, database, cleaning, and URL-inspection behavior is provided by modules and functions.",
    360: "The diagram shows how Flask routes invoke these services and how shared database and notification utilities support the request workflow.",
    365: "The sequence diagram illustrates communication among the user, React application or extension, Flask API, detection service, external intelligence provider, and database during a scan.",
    366: "The frontend sends JSON or multipart data to Flask. Flask validates the request and invokes the phishing, URL, or file service. Optional AbuseIPDB or VirusTotal calls are made only by the relevant module. The result is stored as a scan summary and returned as JSON.",
    367: "Threat-email delivery is attempted after selected high-risk results. Notification failure is logged and does not change the successful scan response.",
    374: "The final implementation includes binary phishing email detection, visible embedded-link analysis, TXT/EML/PDF/DOCX text input, rule-based URL analysis, VirusTotal file lookup and live upload, scan-summary storage, browser-generated PDF reports, prototype authentication and administration, SMTP notifications, a Gmail Chrome extension, public dashboard statistics, and live deployment. PostgreSQL is used live and SQLite is the local fallback.",
    375: "The frontend handles interaction and report generation, Flask routes coordinate validation and services, the phishing pipeline performs ML classification, URLAnalyzer performs deterministic scoring, MalwareScanner communicates with VirusTotal, and four relational tables store users and scan summaries.",
    379: "The phishing email module analyzes text using Natural Language Processing and Machine Learning. Its Logistic Regression model is binary and predicts ham or spam; the interface maps these labels to Safe or Phishing. Word-level TF-IDF captures phrases, while character-level TF-IDF captures obfuscation and character patterns.",
    392: "12. If the prediction is spam with confidence above 0.70, the notification module may attempt a phishing alert.",
    394: "• Binary classification result (ham/spam, displayed as Safe/Phishing)",
    398: "• Prediction processing time in backend logs",
    405: "3. Up to five unique visible HTTP/HTTPS URLs are collected for analysis.",
    413: "5.1.3 Email and Document File Upload",
    414: "The email page accepts TXT and EML files as decoded text input and also accepts PDF and DOCX documents for text extraction. The extracted text is sent through the same phishing pipeline. The current version does not parse MIME parts, inspect HTML anchor tags, or compare visible link text with hidden href destinations.",
    416: "1. User uploads a TXT, EML, PDF, or DOCX file.",
    417: "2. The backend validates the filename, extension, size, and non-empty content.",
    418: "3. TXT and EML content is decoded using UTF-8 with a Latin-1 fallback.",
    419: "4. PDF text is extracted page by page; DOCX text is collected from non-empty paragraphs.",
    420: "5. Extracted text must contain at least ten readable characters.",
    421: "6. At most the first 5,000 characters of document text are passed to the phishing classifier.",
    422: "7. Visible HTTP/HTTPS URLs present in the extracted text may be analyzed separately.",
    423: "8. The response includes the filename, extraction method, extracted character count, prediction, confidence, threats, tips, and embedded-link results.",
    424: "9. A summary is stored in email_scans using the supplied user-email header or anonymous as fallback.",
    425: "10. Image-only or unreadable documents return a clear extraction error rather than a fabricated prediction.",
    435: "7. The structural score and AbuseIPDB reputation value are combined when reputation data is available.",
    436: "8. Thresholds convert the final score into Safe, Suspicious, or Malicious.",
    437: "9. The response includes the score, confidence, threat indicators, AbuseIPDB status, and safety tips.",
    438: "10. The result is displayed and may trigger an alert only when it is malicious with confidence above 0.60.",
    439: "11. The URL scan summary is saved in url_scans.",
    451: "3. The backend validates the extension, non-empty content, and maximum size of 32 MB.",
    457: "9. The temporary file is deleted after the lookup or live-analysis workflow; the application does not retain the uploaded file.",
    459: "11. An email notification is attempted only when the final verdict is Malicious and at least one engine detected the file.",
    471: "9. The final verdict is displayed and its metadata is saved in file_scans.",
    472: "10. Only a Malicious final verdict triggers an attempted malware email alert.",
    474: "The history module presents the latest scan summaries across email, URL, and file modules. In the current prototype, the backend returns the most recent twenty system-wide scans rather than filtering by the logged-in user.",
    478: "3. The related module inserts a record into email_scans, url_scans, or file_scans.",
    479: "4. Each record stores user_email, module-specific summary and result fields, and scanned_at.",
    480: "5. PDF report files and report paths are not stored in the database.",
    481: "6. The history page requests the latest scan records.",
    484: "9. The current history table does not provide record-detail or report-download actions.",
    486: "The report feature generates a PDF from the currently displayed result in the user browser using html2canvas and jsPDF. The generated file is downloaded locally and is not saved by the backend [30], [31].",
    491: "4. The captured result contains the scan type, verdict, confidence, indicators, and recommendations visible on screen.",
    494: "7. No report record or report path is inserted into the database.",
    496: "The current admin panel is a read-only monitoring interface. It displays system statistics, registered users, recent activity, and the latest fifty scans; it does not edit users, settings, logs, or model files.",
    498: "1. A user logs in with an email configured in the ADMIN_EMAILS environment variable.",
    499: "2. The prototype backend compares the X-User-Email request header with that configured list.",
    501: "4. The Overview tab shows users, scan counts, detected threats, and recent activity.",
    502: "5. The Users tab shows registration dates and scan counts but provides no modification controls.",
    503: "6. The Scans tab shows recent email, URL, and file scan summaries.",
    504: "7. Statistics are calculated from current database records.",
    505: "8. Signed tokens and server-enforced role records are future security work.",
    507: "The notification module sends welcome messages after registration and attempts alerts for configured high-risk scan outcomes. Missing SMTP configuration, anonymous requests, or delivery failures do not interrupt the scan.",
    511: "3. Alerts are attempted for spam predictions above 0.70, malicious URLs above 0.60 confidence, and Malicious file verdicts with at least one detection.",
    514: "6. Delivery success or failure is written to application logs; there is no notification database table.",
    522: "4. The content script reads the largest visible Gmail message-text block.",
    523: "5. The extension submits up to 5,000 characters to the live /predict endpoint.",
    528: "10. Extension requests do not currently include an authenticated user header and are stored as anonymous if database saving succeeds.",
    532: "1. The Flask backend is deployed on Render.",
    535: "4. The React frontend is deployed on Vercel.",
    538: "7. PostgreSQL is configured for persistent deployed storage, while SQLite remains available locally.",
    543: "CyberSentinel AI integrates VirusTotal for file reputation and live analysis, AbuseIPDB for resolved-IP reputation, and SMTP for welcome and high-risk alert email. API credentials remain in backend environment variables and are not embedded in frontend or extension code [14], [15].",
    547: "These integrations complement the local phishing model and deterministic URL rules. External failures are handled with error, unavailable, pending, or fallback responses instead of being presented as successful threat-intelligence checks.",
    552: "The login and registration interface supports account creation and prototype login state. The current frontend stores a dummy token, email address, and login flag in localStorage; production-grade signed sessions and role records are future work.",
    556: "• Configured admin-email check for the prototype admin view",
    557: "• Client-side login-state handling",
    565: "The dashboard links to the three scanning modules and displays database-backed email-scan and threat counts from the public /api/stats endpoint, refreshed every thirty seconds.",
    571: "• Database-backed system statistics",
    572: "• Access to scan-history and admin-monitoring views",
    581: "The email scan page accepts pasted text, TXT/EML email files, and PDF/DOCX documents. The binary model result is displayed as Safe or Phishing with confidence, indicators, recommendations, and any visible embedded-link results.",
    587: "• TXT/EML upload and PDF/DOCX text extraction",
    594: "5.3.4 Email Result and Embedded URL Interface",
    595: "The result section displays the Safe or Phishing verdict, confidence, threat indicators, safety recommendations, and separately analyzed visible HTTP/HTTPS URLs when any are found. Hidden HTML href mismatch detection is not implemented in the current version.",
    597: "• Email verdict and confidence",
    598: "• Visible extracted URL, when present",
    599: "• Individual URL verdict and risk score",
    600: "• Threat indicators and safety recommendations",
    601: "• Scan-again and PDF-download actions",
    609: "The URL scan page applies ten deterministic structural checks and, when available, AbuseIPDB reputation for the resolved IP address.",
    615: "• AbuseIPDB availability and abuse score",
    622: "The file scan page validates supported extensions and a 32 MB limit, creates a temporary backend file, and uses VirusTotal hash lookup or live upload. The temporary file is deleted, but unknown files may be transmitted to VirusTotal for analysis.",
    630: "• Temporary-file deletion status and clear external-upload disclosure",
    636: "The history page displays the latest twenty scan summaries across all modules. The current backend response is system-wide and the table does not provide record-detail or report-download actions.",
    642: "• Module-specific confidence where available",
    643: "• System-wide latest-record limitation",
    649: "Each result page can capture the displayed result and download it as a PDF through client-side libraries. Reports are not persisted by the backend.",
    657: "• Local browser download",
    663: "The admin panel provides read-only monitoring for configured administrator emails. It exposes Overview, Users, and Scans tabs backed by database queries.",
    668: "• View registered users and per-module scan counts",
    669: "• View recent scan summaries",
    670: "• Refresh current database statistics",
    671: "• No user-editing, log-management, settings, or model controls",
    676: "Email output is generated for welcome messages and selected high-risk results when SMTP credentials are configured. The example shown is a high-confidence phishing alert.",
    688: "The Gmail extension reads the visible text of an opened Gmail message and sends it to the live phishing endpoint. The popup displays a binary verdict, confidence, threat indicators, and any visible URLs returned by the backend.",
    699: "This chapter described the implemented ML email classifier, document-text extraction, visible embedded-link analysis, deterministic URL scoring, VirusTotal workflow, PostgreSQL/SQLite storage, client-side reports, prototype administration, notifications, Gmail extension, and deployed architecture without attributing unimplemented features to the system.",
    704: "The final version includes binary phishing detection, visible embedded-link scanning, TXT/EML/PDF/DOCX input, rule-based URL analysis, VirusTotal file lookup and live upload, latest scan history, client-side PDF reports, read-only admin monitoring, configured email notifications, a Gmail extension, and live deployment.",
    712: "• Email and document file upload",
    740: "Unit Testing 3: Email and Document File Upload",
    741: "Testing Objective: To verify TXT/EML decoding, PDF/DOCX text extraction, and unsupported or unreadable-file handling.",
    756: "Testing Objective: To verify account lookup, prototype login responses, database inserts, and latest-history retrieval.",
    761: "Testing Objective: To verify read-only admin monitoring and configured notification behavior.",
    778: "Python integration scripts and Postman requests were used to exercise core backend endpoints, validation, and response structures. These scripts are useful regression checks but are not a comprehensive pytest or unittest suite.",
    779: "Additional verification used browser developer tools, database inspection, Chrome extension developer mode, and deployment logs. The React repository still contains the default Create React App test and does not provide substantive automated frontend coverage.",
    789: "6.2.3 Model and Performance Evaluation",
    790: "The saved binary Logistic Regression model was evaluated by recreating the stratified 80/20 split of phishing_clean_balanced.csv with random_state=42. On 1,333 held-out records it achieved 98.05% accuracy. Spam precision was 99.02%, recall 94.82%, and F1-score 96.88%; the confusion matrix contained 904 true ham, 4 ham classified as spam, 403 true spam, and 22 spam classified as ham. API response times remain dependent on network and external-service availability.",
    795: "Security testing covered sanitization, validation, rate limiting, payload controls, environment-based API secrets, CORS, and temporary-file deletion. Authentication and administrator authorization remain prototype controls and are identified as partial rather than production-grade protections [17], [18].",
    800: "Testing supports the implemented scanning, storage, reporting, notification, extension, and deployment workflows. False hidden-link, PhishTank, user-management, audit-log, and report-storage claims have been removed. Prototype authentication, global history visibility, and limited automated test coverage remain documented limitations.",
    801: "The final system is suitable for an academic demonstration of a hybrid cybersecurity toolkit. Production use would require stronger identity controls, per-user data isolation, formal automated tests, operational monitoring, and additional privacy review.",
    805: "CyberSentinel AI provides one deployed interface for phishing-email classification, visible URL inspection, and VirusTotal-assisted file analysis. Supporting features include document-text extraction, latest scan history, browser-generated reports, read-only admin monitoring, configured notifications, and a Gmail extension.",
    807: "CyberSentinel AI achieves its main objective as a unified and understandable cybersecurity prototype. It combines different techniques according to the threat type instead of inaccurately describing every module as machine learning.",
    808: "The phishing module uses hybrid word and character TF-IDF with binary Logistic Regression. The model labels are ham and spam and are displayed as Safe and Phishing. On the recreated held-out split of the final balanced dataset, the saved model achieved 98.05% accuracy, with 99.02% spam precision and 94.82% spam recall.",
    809: "The URL module analyzes length, keywords, IP-based hosts, scheme, shorteners, top-level domains, subdomains, path length, and related structural indicators. AbuseIPDB adds optional resolved-IP reputation; PhishTank is not integrated in the final system.",
    811: "The email workflow extracts and analyzes visible HTTP/HTTPS URLs. TXT and EML files are treated as decoded text, while PDF and DOCX files use text extraction. Structured MIME parsing, HTML anchor extraction, and visible-text versus href mismatch detection remain future work.",
    812: "The final version also includes prototype authentication, PostgreSQL-backed scan summaries, browser-generated reports, read-only admin monitoring, application logging, configured email alerts, a Gmail extension, and Render/Vercel deployment. Reports and application logs are not stored as database entities.",
    813: "Manual and script-based verification covered the implemented workflows. The current limitations include prototype token handling, header-based admin checks, global rather than per-user history, and limited formal automated test coverage.",
    816: "The deployed prototype is functional, but the following improvements are required for stronger accuracy, privacy, authorization, scalability, and production readiness.",
    822: "The Gmail extension can be expanded into a browser-wide security assistant. Future email work should also add standards-based MIME parsing, HTML anchor extraction, and visible-text versus destination-domain mismatch detection.",
    832: "Authentication should replace the dummy client token and unsalted SHA-256 password hashing with signed server sessions or JWTs and a slow salted password hash such as Argon2 or bcrypt. Administrator roles must be verified server-side, and scan-history queries should be restricted to the authenticated user. Two-factor authentication, session expiry, and stronger password policies can then be added.",
    838: "CyberSentinel AI has been developed as a deployed AI-assisted prototype with binary phishing detection, visible embedded-link analysis, rule-based URL scoring, VirusTotal-assisted file analysis, scan summaries, browser reports, read-only admin monitoring, notifications, and a Gmail extension. Future work includes structured EML analysis, stronger authentication and privacy, formal automated testing, larger multilingual datasets, mobile access, behavioral malware analysis, and scalable monitoring.",
}


CAPTIONS = {
    362: ": Class and Module Diagram",
    592: ": Email Phishing Scan Interface",
    605: ": Email Phishing Scan Result",
    618: ": URL Threat Analysis Interface",
    659: ": PDF Report Generation Option",
    743: ": Email and Document File Upload Unit Test Case",
    782: ": Tools Used for Automated Testing",
    792: ": Model and Performance Results",
}


def patch_tables(document):
    tables = document.tables
    set_cell(tables[2].cell(3, 1), "User")
    set_cell(tables[2].cell(4, 1), "Backend API, rule-based URL analyzer, AbuseIPDB")
    set_cell(tables[4].cell(4, 1), "System shall analyze URLs using rule-based scoring and optional AbuseIPDB reputation.")

    # Remove the unimplemented PhishTank integration row.
    delete_row(tables[6], 3)
    set_row(tables[6], 3, [
        "SMTP / Email Service",
        "Configured SMTP service for welcome messages and selected high-risk alerts.",
        "Used after signup and after configured phishing, URL, or malware outcomes.",
        "notification_service.py, send_welcome_email(), send_threat_alert()",
    ])

    set_row(tables[7], 2, ["2", "Register a new user account", "Email and password", "User account is created successfully", "Pass"])
    set_row(tables[7], 5, ["5", "Submit phishing email text", "JazzCash-style phishing message", "System returns spam and the interface displays Phishing", "Pass"])
    set_row(tables[7], 8, ["8", "Upload a PDF or DOCX containing email text", "Readable document", "Text is extracted and classified", "Pass"])
    set_row(tables[7], 14, ["14", "View scan history", "Existing scan records", "Latest system-wide scan summaries are displayed", "Pass"])
    set_row(tables[7], 17, ["17", "Trigger configured risk notification", "High-confidence phishing, malicious URL, or malicious file", "SMTP alert is attempted without blocking the scan", "Pass"])

    set_row(tables[8], 1, ["1", "Test phishing email classification", "Your JazzCash account is suspended. Verify now.", "Email classified as spam / displayed as Phishing", "Pass"])
    set_row(tables[9], 2, ["2", "Extract multiple URLs", "Email contains several URLs", "Up to five unique URLs are analyzed", "Pass"])

    replacement = [
        ["No.", "Test Case / Test Script", "Attribute and Value", "Expected Result", "Result"],
        ["1", "Upload TXT email file", "Readable UTF-8 text", "Content decoded and classified", "Pass"],
        ["2", "Upload EML file", "Readable email source", "File treated as decoded text and classified", "Pass"],
        ["3", "Upload text-based PDF", "PDF containing email text", "Page text extracted and classified", "Pass"],
        ["4", "Upload DOCX document", "DOCX containing email text", "Paragraph text extracted and classified", "Pass"],
        ["5", "Upload unreadable or unsupported file", "Image-only/empty/unsupported input", "Clear validation or extraction error", "Pass"],
    ]
    for index, values in enumerate(replacement):
        set_row(tables[10], index, values)

    set_row(tables[13], 3, ["3", "Login valid user", "Correct credentials", "Success response and prototype token returned", "Pass"])
    set_row(tables[13], 8, ["8", "Retrieve scan history", "History request", "Latest system-wide records returned", "Pass"])
    set_row(tables[14], 3, ["3", "View recent scans", "Admin Scans tab", "Latest scan summaries displayed", "Pass"])
    set_row(tables[14], 5, ["5", "Trigger threat notification", "Configured high-risk result", "Email is attempted when SMTP is configured", "Pass"])

    set_row(tables[15], 6, ["6", "Email/document file upload", "Upload TXT, EML, PDF, or DOCX", "Readable text is classified", "Pass"])
    set_row(tables[15], 12, ["12", "Scan history", "Open history page", "Latest system-wide scans are displayed", "Pass"])
    set_row(tables[15], 14, ["14", "Admin monitoring", "Configured admin opens panel", "Statistics, users, and scans load", "Pass"])
    set_row(tables[15], 15, ["15", "Email notification", "Configured high-risk result", "Alert is attempted without blocking scan", "Pass"])

    set_row(tables[16], 4, ["4", "Document extractor + phishing model", "Upload PDF or DOCX", "Readable text is extracted and classified", "Pass"])
    set_row(tables[16], 9, ["9", "Frontend + scan history API", "Open history page", "Latest records across modules are displayed", "Pass"])
    set_row(tables[16], 10, ["10", "Client report generator + scan result", "Download PDF", "Current result is captured and downloaded locally", "Pass"])
    set_row(tables[16], 12, ["12", "Admin panel + database", "Open admin tabs", "Statistics, users, and scans load", "Pass"])
    set_row(tables[16], 13, ["13", "Gmail extension + live backend", "Scan visible Gmail text", "Binary prediction returned to popup", "Pass"])

    set_row(tables[17], 3, ["Python integration scripts", "Custom request scripts for endpoint and validation checks.", "Core backend endpoints", "Useful regression checks; not a formal pytest suite"])
    set_row(tables[17], 5, ["Database inspection", "SQLite Browser locally and deployment database queries.", "users and three scan tables", "Records stored and retrieved"])
    set_row(tables[17], 7, ["Render logs", "Deployment logs used to inspect backend behavior.", "Live Flask backend", "Backend reachable and errors inspectable"])
    set_row(tables[17], 8, ["Vercel logs", "Frontend deployment logs and network inspection.", "Live React frontend", "Frontend connected to backend"])

    set_row(tables[18], 6, ["6", "Email File Upload API", "Valid TXT/EML file", "Decoded text and prediction returned", "Pass"])
    set_row(tables[18], 12, ["12", "Scan History API", "GET request", "Latest system-wide scan records returned", "Pass"])

    model_row = tables[19].add_row()
    for cell, text in zip(model_row.cells, [
        "10", "Final phishing model evaluation", "Held-out 20% split", "98.05% accuracy; spam F1 96.88%", "Pass"
    ]):
        set_cell(cell, text)

    set_row(tables[20], 7, ["7", "Password storage", "Inspect stored value", "Plain password is not stored", "Partial - SHA-256 is unsalted"])
    set_row(tables[20], 8, ["8", "Prototype admin access", "Normal and configured-admin headers", "Configured list is checked", "Partial - header is client supplied"])


def patch_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for name, size, before, after in [
        ("Heading 1", 18, 12, 6),
        ("Heading 2", 14, 10, 5),
        ("Heading 3", 13, 8, 4),
    ]:
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(6)


def replace_media(docx_path):
    replacements = {
        f"word/media/image{number}.png": DIAGRAMS / f"image{number}.png"
        for number in range(3, 9)
    }
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=docx_path.parent) as handle:
        temp_path = Path(handle.name)
    try:
        with ZipFile(docx_path, "r") as source_zip, ZipFile(temp_path, "w", ZIP_DEFLATED) as target_zip:
            for info in source_zip.infolist():
                data = replacements[info.filename].read_bytes() if info.filename in replacements else source_zip.read(info.filename)
                target_zip.writestr(info, data)
        temp_path.replace(docx_path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def main():
    document = Document(SOURCE)
    if len(document.paragraphs) != 882 or len(document.tables) != 21:
        raise RuntimeError("Unexpected source report structure; refusing an unsafe positional edit.")

    document.core_properties.title = "CyberSentinel AI - Final Year Project Report"
    document.core_properties.author = "Mohavia Arif"
    document.core_properties.last_modified_by = "Mohavia Arif"

    for index, text in PARAGRAPHS.items():
        set_paragraph_text(document.paragraphs[index], text)
    for index, suffix in CAPTIONS.items():
        replace_caption_suffix(document.paragraphs[index], suffix)

    patch_tables(document)
    patch_styles(document)
    add_front_matter_page_breaks(document)
    add_front_body_section_break(document)
    document.save(OUTPUT)

    # Reload so python-docx resolves both newly created sections.
    document = Document(OUTPUT)
    if len(document.sections) != 2:
        raise RuntimeError(f"Expected 2 sections after front-matter split, found {len(document.sections)}")
    front, body = document.sections
    front.different_first_page_header_footer = True
    body.different_first_page_header_footer = False
    front.header.is_linked_to_previous = False
    front.footer.is_linked_to_previous = False
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    front.header_distance = Inches(0.35)
    front.footer_distance = Inches(0.35)
    body.header_distance = Inches(0.35)
    body.footer_distance = Inches(0.35)
    set_page_numbering(front, "lowerRoman", 1)
    set_page_numbering(body, "decimal", 1)
    clear_story(front.header)
    clear_story(front.first_page_header)
    clear_story(front.first_page_footer)
    build_footer(front)
    build_header(body, "CyberSentinel AI - Intelligent Cybersecurity Toolkit")
    build_footer(body)
    document.save(OUTPUT)

    replace_media(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
