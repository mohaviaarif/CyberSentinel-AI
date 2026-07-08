from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "CyberSentinel_Final_Report_Final.docx"
ORIGINAL = ROOT / "CyberSentinel_Final_Report_Original.docx"
OUTPUT = ROOT / "CyberSentinel_Final_Report_Tables_Corrected_Diagrams_Restored.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(9)
HEADER_FILL = "D9E2F3"
LABEL_FILL = "E7E6E6"


def set_cell_text(cell, text):
    cell.text = text


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def prepare_row(row, repeat_header=False):
    tr_pr = row._tr.get_or_add_trPr()
    for height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(height)
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    if repeat_header:
        header = tr_pr.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            tr_pr.append(header)
        header.set(qn("w:val"), "true")


def format_run(run, bold=False):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)


def format_cell(cell, alignment, bold=False, fill="FFFFFF"):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            format_run(run, bold=bold)


def normalize_table(table, table_index):
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        prepare_row(row, repeat_header=row_index == 0)
        for column_index, cell in enumerate(row.cells):
            if row_index == 0:
                format_cell(
                    cell,
                    WD_ALIGN_PARAGRAPH.CENTER,
                    bold=True,
                    fill=HEADER_FILL,
                )
                continue

            is_label_table = table_index <= 5
            label_cell = is_label_table and column_index == 0
            fill = LABEL_FILL if label_cell else "FFFFFF"
            bold = label_cell

            if table_index == 0:
                alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
            elif table_index in (1, 2, 3):
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif table_index in (4, 5):
                alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
            elif table_index == 17:
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif len(table.columns) == 5 and column_index in (0, 4):
                alignment = WD_ALIGN_PARAGRAPH.CENTER
                if column_index == 4:
                    bold = True
            else:
                alignment = WD_ALIGN_PARAGRAPH.LEFT

            format_cell(cell, alignment, bold=bold, fill=fill)


def correct_table_wording(document):
    tables = document.tables

    set_cell_text(tables[0].cell(3, 1), "Comma-Separated Values")
    set_cell_text(tables[0].cell(9, 1), "Term Frequency-Inverse Document Frequency")

    use_cases = [
        [
            (0, "Field", "Description"),
            (1, "Use Case ID", "UC-01"),
            (2, "Use Case Name", "Scan Email for Phishing"),
            (3, "Primary Actor", "User"),
            (4, "Secondary Actors", "Backend API, NLP module, and ML model"),
            (5, "Description", "The user submits email text for phishing detection"),
            (6, "Trigger", 'The user clicks "Analyze Email"'),
            (7, "Preconditions", "Internet connection, running backend, and loaded model"),
            (8, "Postconditions", "The prediction is displayed and the scan is stored"),
        ],
        [
            (0, "Field", "Description"),
            (1, "Use Case ID", "UC-02"),
            (2, "Use Case Name", "Analyze URL for Threats"),
            (3, "Primary Actor", "User"),
            (4, "Secondary Actors", "Backend API, rule-based URL analyzer, and AbuseIPDB"),
            (5, "Description", "The user submits a URL for threat analysis"),
            (6, "Trigger", 'The user clicks "Analyze URL"'),
            (7, "Preconditions", "Internet connection and running backend"),
            (8, "Postconditions", "The risk result is displayed and the scan is stored"),
        ],
        [
            (0, "Field", "Description"),
            (1, "Use Case ID", "UC-03"),
            (2, "Use Case Name", "Scan File for Malware"),
            (3, "Primary Actor", "User"),
            (4, "Secondary Actors", "Backend API and VirusTotal API"),
            (5, "Description", "The user uploads a file for malware scanning"),
            (6, "Trigger", "The user selects and uploads a file"),
            (7, "Preconditions", "Supported file, internet connection, and VirusTotal API access"),
            (8, "Postconditions", "The malware verdict is displayed and the scan is stored"),
        ],
    ]
    for table_index, rows in zip((1, 2, 3), use_cases):
        for row_index, left, right in rows:
            set_cell_text(tables[table_index].cell(row_index, 0), left)
            set_cell_text(tables[table_index].cell(row_index, 1), right)

    functional_requirements = [
        "The system shall allow users to submit email text.",
        "The system shall classify email using NLP and machine learning.",
        "The system shall allow users to submit URLs.",
        "The system shall analyze URLs using rule-based scoring and optional AbuseIPDB reputation.",
        "The system shall allow users to upload files.",
        "The system shall integrate with the VirusTotal API.",
        "The system shall display results using clear color indicators.",
        "The system shall store scan-history records.",
        "The system shall generate downloadable reports.",
        "The system shall provide read-only administrator monitoring.",
    ]
    for row_index, requirement in enumerate(functional_requirements, 1):
        set_cell_text(tables[4].cell(row_index, 1), requirement)

    non_functional_requirements = [
        "The system shall provide a simple interface for non-technical users.",
        "The system shall return results within a few seconds when external services respond normally.",
        "The system shall transmit production traffic over HTTPS.",
        "The system shall handle external-service failures gracefully.",
        "The system shall support modular updates.",
    ]
    for row_index, requirement in enumerate(non_functional_requirements, 1):
        set_cell_text(tables[5].cell(row_index, 1), requirement)

    api_headers = ["API / Service", "Description", "Purpose", "Implementation"]
    for column_index, header in enumerate(api_headers):
        set_cell_text(tables[6].cell(0, column_index), header)
    set_cell_text(
        tables[6].cell(5, 3),
        "manifest.json, content.js, and popup.js",
    )

    for table_index in range(8, 15):
        set_cell_text(tables[table_index].cell(0, 2), "Input / Attribute")

    set_cell_text(tables[7].cell(15, 1), "Open the admin panel as a configured administrator")
    set_cell_text(tables[7].cell(15, 2), "Configured administrator email")
    set_cell_text(tables[7].cell(15, 3), "Admin panel is displayed")

    tool_names = {
        2: "Python Integration Scripts",
        4: "Database Inspection",
        6: "Render Logs",
        7: "Vercel Logs",
    }
    for row_index, value in tool_names.items():
        set_cell_text(tables[17].cell(row_index, 0), value)
    set_cell_text(
        tables[17].cell(4, 2),
        "Users table and the three scan tables",
    )
    set_cell_text(
        tables[18].cell(1, 3),
        "Success response with prototype token",
    )


def replace_caption_suffix(paragraph, suffix):
    if paragraph.runs:
        paragraph.runs[-1].text = suffix


def correct_table_captions(document):
    replacements = {
        "Table 3. 1:": "Use Case Description - Scan Email for Phishing",
        "Table 3. 2:": "Use Case Description - Scan URL for Threats",
        "Table 3. 3:": ": Use Case Description - Scan File for Malware",
        "Table 3. 4:": ": Functional Requirements",
    }
    for paragraph in document.paragraphs:
        if paragraph.style.name.lower() != "caption":
            continue
        text = paragraph.text.strip()
        for prefix, suffix in replacements.items():
            if text.startswith(prefix):
                replace_caption_suffix(paragraph, suffix)
                break


def correct_table_list_entries(document):
    replacements = {
        "Table 3. 1: USE CASE": "Table 3. 1: Use Case Description - Scan Email for Phishing",
        "Table 3. 2: Use Case Description – Scan URL for Threats": "Table 3. 2: Use Case Description - Scan URL for Threats",
        "Table 3. 3: Use Case Description – Malware File Detection": "Table 3. 3: Use Case Description - Scan File for Malware",
        "Table 3. 4:  Functional Requirements": "Table 3. 4: Functional Requirements",
    }
    for paragraph in document.paragraphs:
        if paragraph.style.name.lower() != "table of figures":
            continue
        for text_node in paragraph._p.iter(qn("w:t")):
            if text_node.text in replacements:
                text_node.text = replacements[text_node.text]


def restore_original_diagrams(docx_path):
    with ZipFile(ORIGINAL, "r") as original_zip:
        original_images = {
            f"word/media/image{index}.png": original_zip.read(
                f"word/media/image{index}.png"
            )
            for index in range(3, 9)
        }

    with ZipFile(docx_path, "r") as source_zip, NamedTemporaryFile(
        suffix=".docx", delete=False, dir=ROOT
    ) as temporary_file:
        temporary_path = Path(temporary_file.name)
        with ZipFile(temporary_file, "w", ZIP_DEFLATED) as output_zip:
            for item in source_zip.infolist():
                data = original_images.get(item.filename, source_zip.read(item.filename))
                output_zip.writestr(item, data)

    temporary_path.replace(docx_path)


def main():
    document = Document(SOURCE)
    correct_table_wording(document)
    correct_table_captions(document)
    correct_table_list_entries(document)
    for table_index, table in enumerate(document.tables):
        normalize_table(table, table_index)
    document.save(OUTPUT)
    restore_original_diagrams(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
